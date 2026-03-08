"""
Generate Speaker Notes Word Documents
Converts all speaker-notes .md files to professionally styled .docx files.
Uses python-docx to create a brand-styled reference template, then Pandoc for conversion.

Usage: python generate_speaker_notes_docx.py
"""

import os
import subprocess
from docx import Document
from docx.shared import Pt, RGBColor, Cm, Emu
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
REFERENCE_DOC = os.path.join(SCRIPT_DIR, "_speaker-notes-reference.docx")

# BKA brand colors
BKA_BLUE = RGBColor(0x00, 0x5C, 0xB9)
DARK_GRAY = RGBColor(0x33, 0x33, 0x33)
MID_GRAY = RGBColor(0x66, 0x66, 0x66)
LIGHT_GRAY = RGBColor(0x99, 0x99, 0x99)

# All speaker notes files: (source_md, output_docx)
FILES = [
    ("intro-video-speaker-notes.md", "intro-video-speaker-notes.docx"),
    ("ch01-intro/speaker-notes.md", "ch01-intro/ch01-speaker-notes.docx"),
    ("ch02-qualitative-models/speaker-notes.md", "ch02-qualitative-models/ch02-speaker-notes.docx"),
    ("ch03-interview-case/speaker-notes.md", "ch03-interview-case/ch03-speaker-notes.docx"),
    ("ch04-quantitative-models/speaker-notes.md", "ch04-quantitative-models/ch04-speaker-notes.docx"),
    ("ch05-situational-analysis/speaker-notes.md", "ch05-situational-analysis/ch05-speaker-notes.docx"),
    ("ch06-business-financials/speaker-notes.md", "ch06-business-financials/ch06-speaker-notes.docx"),
    ("ch07-single-agent/speaker-notes.md", "ch07-single-agent/ch07-speaker-notes.docx"),
    ("ch08-resource-allocation/speaker-notes.md", "ch08-resource-allocation/ch08-speaker-notes.docx"),
    ("ch09-multi-agent/speaker-notes.md", "ch09-multi-agent/ch09-speaker-notes.docx"),
    ("ch10-data-driven/speaker-notes.md", "ch10-data-driven/ch10-speaker-notes.docx"),
]


def set_font(style, name, size_pt, color, bold=False, italic=False):
    """Configure font properties on a style."""
    font = style.font
    font.name = name
    font.size = Pt(size_pt)
    font.color.rgb = color
    font.bold = bold
    font.italic = italic


def add_header_footer(doc):
    """Add header and footer to the reference document."""
    for section in doc.sections:
        # Header: "Inference and Intervention" right-aligned
        header = section.header
        header.is_linked_to_previous = False
        header_para = header.paragraphs[0]
        header_para.text = "Inference and Intervention"
        header_para.alignment = 2  # RIGHT
        run = header_para.runs[0]
        run.font.name = "Calibri"
        run.font.size = Pt(9)
        run.font.color.rgb = LIGHT_GRAY

        # Footer: "BK Advisors" left, page number right
        footer = section.footer
        footer.is_linked_to_previous = False
        footer_para = footer.paragraphs[0]
        footer_para.text = ""

        # Add "BK Advisors" run
        run_left = footer_para.add_run("BK Advisors")
        run_left.font.name = "Calibri"
        run_left.font.size = Pt(9)
        run_left.font.color.rgb = LIGHT_GRAY

        # Add tab and page number field
        run_tab = footer_para.add_run("\t\t")
        run_tab.font.size = Pt(9)

        # Page number field
        fld_char_begin = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
        run_pg = footer_para.add_run()
        run_pg._r.append(fld_char_begin)

        instr_text = parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> PAGE </w:instrText>')
        run_pg2 = footer_para.add_run()
        run_pg2._r.append(instr_text)

        fld_char_end = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
        run_pg3 = footer_para.add_run()
        run_pg3._r.append(fld_char_end)
        run_pg3.font.name = "Calibri"
        run_pg3.font.size = Pt(9)
        run_pg3.font.color.rgb = LIGHT_GRAY


def build_reference_doc():
    """Create a Pandoc reference .docx with BKA brand styles."""
    doc = Document()

    # Page setup
    for section in doc.sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(2.54)
        section.right_margin = Cm(2.54)

    # Normal style
    style = doc.styles["Normal"]
    set_font(style, "Calibri", 11, DARK_GRAY)
    style.paragraph_format.space_after = Pt(6)
    style.paragraph_format.line_spacing = 1.15

    # Heading 1
    h1 = doc.styles["Heading 1"]
    set_font(h1, "Calibri Light", 26, BKA_BLUE, bold=True)
    h1.paragraph_format.space_before = Pt(0)
    h1.paragraph_format.space_after = Pt(4)

    # Heading 2
    h2 = doc.styles["Heading 2"]
    set_font(h2, "Calibri Light", 16, MID_GRAY, italic=True)
    h2.paragraph_format.space_before = Pt(18)
    h2.paragraph_format.space_after = Pt(6)

    # Heading 3
    h3 = doc.styles["Heading 3"]
    set_font(h3, "Calibri Light", 14, BKA_BLUE)
    h3.paragraph_format.space_before = Pt(12)
    h3.paragraph_format.space_after = Pt(4)

    # First Paragraph / Body Text styles (used by Pandoc)
    for style_name in ["Body Text", "First Paragraph"]:
        try:
            s = doc.styles[style_name]
            set_font(s, "Calibri", 11, DARK_GRAY)
            s.paragraph_format.space_after = Pt(6)
            s.paragraph_format.line_spacing = 1.15
        except KeyError:
            pass

    # List Bullet style
    try:
        lb = doc.styles["List Bullet"]
        set_font(lb, "Calibri", 11, DARK_GRAY)
        lb.paragraph_format.space_after = Pt(4)
    except KeyError:
        pass

    # Add header and footer
    add_header_footer(doc)

    # Add a placeholder paragraph (Pandoc needs content to map styles)
    doc.add_heading("Title", level=1)
    doc.add_heading("Subtitle", level=2)
    doc.add_heading("Section", level=3)
    doc.add_paragraph("Body text.")

    doc.save(REFERENCE_DOC)
    print(f"Reference template: {REFERENCE_DOC}")


def find_pandoc():
    """Find a working pandoc command."""
    # Try quarto pandoc first
    try:
        subprocess.run(
            ["quarto", "pandoc", "--version"],
            capture_output=True, check=True
        )
        return ["quarto", "pandoc"]
    except (FileNotFoundError, subprocess.CalledProcessError):
        pass

    # Try pandoc directly
    try:
        subprocess.run(
            ["pandoc", "--version"],
            capture_output=True, check=True
        )
        return ["pandoc"]
    except (FileNotFoundError, subprocess.CalledProcessError):
        pass

    raise RuntimeError("Neither 'quarto pandoc' nor 'pandoc' found on PATH")


def convert_files():
    """Convert all speaker notes .md → .docx using Pandoc."""
    pandoc_cmd = find_pandoc()
    print(f"Using: {' '.join(pandoc_cmd)}")

    count = 0
    for md_rel, docx_rel in FILES:
        md_path = os.path.join(SCRIPT_DIR, md_rel)
        docx_path = os.path.join(SCRIPT_DIR, docx_rel)

        if not os.path.isfile(md_path):
            print(f"  SKIP (not found): {md_rel}")
            continue

        cmd = pandoc_cmd + [
            md_path,
            "-o", docx_path,
            "--reference-doc", REFERENCE_DOC,
            "--from", "markdown",
            "--to", "docx",
        ]

        result = subprocess.run(cmd, capture_output=True, text=True)
        if result.returncode != 0:
            print(f"  ERROR: {md_rel}")
            print(f"         {result.stderr.strip()}")
        else:
            print(f"  OK: {md_rel} -> {docx_rel}")
            count += 1

    return count


if __name__ == "__main__":
    print("Building reference template...")
    build_reference_doc()
    print()

    print("Converting speaker notes...")
    converted = convert_files()
    print(f"\nDone. Converted {converted} of {len(FILES)} files.")
