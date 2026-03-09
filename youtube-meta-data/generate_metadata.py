"""
Generate YouTube metadata Word documents for each chapter video
plus a course intro video. Run: python generate_metadata.py
"""

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

# ---------------------------------------------------------------------------
# Shared boilerplate
# ---------------------------------------------------------------------------

COURSE_TITLE = "Inference and Intervention: Causal Thinking for Maternal & Newborn Health"
CATEGORY = "Education"

COURSE_FOOTER = (
    'This is part of "{course}" — a free 10-part course teaching causal '
    "reasoning with real-world health examples from Ethiopia, Rwanda, "
    "Kenya, and Tanzania.\n\n"
    "Full course: [LINK]\n"
    "Slides & resources: [LINK]"
)

HASHTAGS = "#CausalInference #GlobalHealth #MaternalHealth #NewbornHealth #DataScience #PublicHealth"

BASE_TAGS = [
    "causal inference",
    "maternal health",
    "newborn health",
    "global health",
    "public health",
    "MNH",
    "causal models",
    "DAG",
    "directed acyclic graph",
    "Ethiopia",
    "Rwanda",
    "Kenya",
    "Tanzania",
    "health data",
    "data science for health",
]

# ---------------------------------------------------------------------------
# Per-video metadata
# ---------------------------------------------------------------------------

VIDEOS = [
    # 0 — Course Intro
    {
        "filename": "00-course-intro.docx",
        "title": "Causal Thinking for Maternal & Newborn Health — Free Course Introduction",
        "description": (
            "Why do some health programs succeed while others fail? The answer often comes down to "
            "understanding cause and effect — not just correlation.\n\n"
            "Welcome to Inference and Intervention: Causal Thinking for Maternal & Newborn Health. "
            "In this free 10-part video course, you will learn how to build causal models, put "
            "numbers on them, and use them to make better decisions about where to invest resources "
            "to save the most lives.\n\n"
            "Every example is drawn from real, publicly available data on maternal and newborn "
            "health challenges in Ethiopia, Rwanda, Kenya, and Tanzania. No statistics background "
            "is needed — just curiosity about how the world works.\n\n"
            "In this video, you will learn:\n"
            "- What this course covers and who it is for\n"
            "- The four-part structure: Foundations, Adding Numbers, Making Decisions, Learning from Data\n"
            "- How causal reasoning applies to real health programs\n"
            "- What tools and resources you will use\n\n"
            "TIMESTAMPS\n"
            "00:00 Welcome\n"
            "00:00 Why causal thinking matters\n"
            "00:00 Course overview — 10 chapters in 4 parts\n"
            "00:00 The four countries: Ethiopia, Rwanda, Kenya, Tanzania\n"
            "00:00 Prerequisites and how to follow along\n"
            "00:00 Getting started\n\n"
        ),
        "tags": BASE_TAGS + [
            "course introduction",
            "free course",
            "causal reasoning",
            "health program evaluation",
            "online course",
        ],
        "thumbnail": "CAUSAL THINKING\nFOR HEALTH\nFree Course",
    },
    # 1 — Introduction to Causal Analysis
    {
        "filename": "01-introduction-to-causal-analysis.docx",
        "title": "Chapter 1: Introduction to Causal Analysis | Correlation Is Not Causation",
        "description": (
            "Ice cream sales and drowning rates both rise in summer — but ice cream does not cause "
            "drowning. This is the most important lesson in data analysis: correlation is not causation.\n\n"
            "In Chapter 1, we introduce the three most common causal thinking traps — confounding, "
            "reverse causation, and selection bias — and show how causal diagrams (DAGs) help you "
            "avoid them. Every example uses real maternal and newborn health scenarios from "
            "Ethiopia, Rwanda, Kenya, and Tanzania.\n\n"
            "In this video, you will learn:\n"
            "- The difference between situational assessment and managerial intervention\n"
            "- Why correlation does not equal causation — with everyday examples\n"
            "- Three causal thinking traps: confounding, reverse causation, selection bias\n"
            "- The basics of causal models: nodes, arrows, and DAGs\n"
            "- Why causal reasoning matters for health programs\n\n"
            "TIMESTAMPS\n"
            "00:00 Introduction\n"
            "00:00 Correlation vs. causation\n"
            "00:00 Trap 1 — Confounding\n"
            "00:00 Trap 2 — Reverse causation\n"
            "00:00 Trap 3 — Selection bias\n"
            "00:00 Causal diagrams (DAGs) — the basics\n"
            "00:00 MNH examples\n"
            "00:00 Key takeaways\n\n"
        ),
        "tags": BASE_TAGS + [
            "correlation vs causation",
            "confounding",
            "reverse causation",
            "selection bias",
            "causal diagrams",
            "introduction to causal inference",
        ],
        "thumbnail": "CHAPTER 1\nCorrelation\n≠ Causation",
    },
    # 2 — Qualitative Causal Models
    {
        "filename": "02-qualitative-causal-models.docx",
        "title": "Chapter 2: Qualitative Causal Models | Chains, Forks, and Colliders",
        "description": (
            "Every causal diagram is built from just three patterns: chains, forks, and colliders. "
            "Once you can spot them, you can trace how information flows — and where it gets blocked.\n\n"
            "In Chapter 2, we learn how to build qualitative causal models with nodes and signed "
            "directed links ([+] and [-]). We walk through d-separation rules that tell you which "
            "variables carry information about each other, and build a Rwanda MNH qualitative DAG "
            "as a worked example.\n\n"
            "In this video, you will learn:\n"
            "- Three types of nodes and how to label them\n"
            "- Signed directed links: positive [+] and negative [-] effects\n"
            "- The three triplet structures: chain, fork, and collider\n"
            "- d-separation rules and how information flows through a DAG\n"
            "- A qualitative DAG for Rwanda's MNH system\n\n"
            "TIMESTAMPS\n"
            "00:00 Introduction\n"
            "00:00 Node types\n"
            "00:00 Signed directed links\n"
            "00:00 Chains — how effects pass through\n"
            "00:00 Forks — common causes\n"
            "00:00 Colliders — information blockers\n"
            "00:00 d-separation rules\n"
            "00:00 Rwanda MNH DAG walkthrough\n"
            "00:00 Key takeaways\n\n"
        ),
        "tags": BASE_TAGS + [
            "qualitative models",
            "chains forks colliders",
            "d-separation",
            "information flow",
            "signed links",
            "Rwanda health",
        ],
        "thumbnail": "CHAPTER 2\nChains, Forks\n& Colliders",
    },
    # 3 — The MNH Diagnostic Case Study
    {
        "filename": "03-mnh-diagnostic-case-study.docx",
        "title": "Chapter 3: Building a Causal Diagram from Interviews | MNH Case Study",
        "description": (
            "How do you go from interviews and program knowledge to a causal diagram? "
            "In this chapter, we walk through the process step by step — starting with a "
            "simple 3-node model and expanding it systematically.\n\n"
            "Using Kenya's maternal and newborn health challenges as our case study, we "
            "practice iterative model-building: form a hypothesis, gather evidence through "
            "interviews, and revise the diagram. You will learn to distinguish root nodes "
            "(leverage points) from intermediate nodes, and to separate supply-side barriers "
            "from demand-side barriers.\n\n"
            "In this video, you will learn:\n"
            "- Iterative model-building: hypothesis → interview → revision\n"
            "- How to expand a 3-4 node model systematically\n"
            "- Root nodes vs. intermediate nodes — where to focus action\n"
            "- Supply-side vs. demand-side barriers in health systems\n"
            "- A worked example using Kenya's MNH data\n\n"
            "TIMESTAMPS\n"
            "00:00 Introduction\n"
            "00:00 Starting with a simple model\n"
            "00:00 The interview process\n"
            "00:00 Adding nodes and arrows\n"
            "00:00 Root nodes vs. intermediate nodes\n"
            "00:00 Supply-side vs. demand-side barriers\n"
            "00:00 Kenya MNH case study\n"
            "00:00 Key takeaways\n\n"
        ),
        "tags": BASE_TAGS + [
            "case study",
            "interview methods",
            "model building",
            "root cause analysis",
            "Kenya health",
            "supply side demand side",
            "health systems",
        ],
        "thumbnail": "CHAPTER 3\nBuilding a DAG\nfrom Interviews",
    },
    # 4 — Quantitative Causal Models
    {
        "filename": "04-quantitative-causal-models.docx",
        "title": "Chapter 4: Quantitative Causal Models | Bayes' Rule Made Simple",
        "description": (
            "Numbers bring a causal diagram to life. In Chapter 4, we learn how to go from "
            "real-world counts to probabilities using natural frequencies — '55 out of 100' "
            "instead of abstract formulas.\n\n"
            "We introduce conditional probability tables (think of them as lookup tables), "
            "walk through Bayes' rule with tree diagrams instead of equations, and build our "
            "first Bayesian network in R using the bnlearn package. All examples use publicly "
            "available MNH data.\n\n"
            "In this video, you will learn:\n"
            "- How to convert real-world counts to probabilities using natural frequencies\n"
            "- Conditional probability tables as simple lookup tables\n"
            "- Bayes' rule — reasoning backwards from outcomes to causes\n"
            "- The Markov Assumption: why each node only depends on its parents\n"
            "- Building a Bayesian network in R with bnlearn\n\n"
            "TIMESTAMPS\n"
            "00:00 Introduction\n"
            "00:00 Natural frequencies — counts, not formulas\n"
            "00:00 Conditional probability tables (lookup tables)\n"
            "00:00 Bayes' rule with tree diagrams\n"
            "00:00 The Markov Assumption\n"
            "00:00 Building a Bayesian network in R\n"
            "00:00 MNH application\n"
            "00:00 Key takeaways\n\n"
        ),
        "tags": BASE_TAGS + [
            "Bayes rule",
            "Bayesian network",
            "conditional probability",
            "natural frequencies",
            "bnlearn",
            "R programming",
            "quantitative models",
        ],
        "thumbnail": "CHAPTER 4\nBayes' Rule\nMade Simple",
    },
    # 5 — Situational Analysis
    {
        "filename": "05-situational-analysis.docx",
        "title": "Chapter 5: Situational Analysis | Updating Beliefs with Evidence",
        "description": (
            "When you learn something new, how should you change your mind? "
            "Situational analysis is about updating your best guess as evidence arrives — "
            "and knowing how information flows (or gets blocked) through a causal diagram.\n\n"
            "In Chapter 5, we trace evidence through chains, forks, and colliders, and "
            'discover the "explaining away" phenomenon — where two competing causes make each '
            "other less likely when the outcome is already known. We apply these ideas to "
            "diagnosing neonatal mortality patterns in Ethiopia.\n\n"
            "In this video, you will learn:\n"
            "- Belief updating — revising your best guess with new data\n"
            "- Marginal probabilities via simple averaging\n"
            "- How evidence flows through chains, forks, and colliders\n"
            '- The "explaining away" phenomenon\n'
            "- Diagnosing Ethiopian neonatal mortality patterns\n\n"
            "TIMESTAMPS\n"
            "00:00 Introduction\n"
            "00:00 Belief updating — changing your mind\n"
            "00:00 Marginal probabilities\n"
            "00:00 Evidence flow through chains\n"
            "00:00 Evidence flow through forks\n"
            "00:00 Evidence flow through colliders\n"
            "00:00 Explaining away\n"
            "00:00 Ethiopia neonatal mortality application\n"
            "00:00 Key takeaways\n\n"
        ),
        "tags": BASE_TAGS + [
            "belief updating",
            "marginal probability",
            "explaining away",
            "evidence flow",
            "Ethiopia health",
            "neonatal mortality",
            "situational analysis",
        ],
        "thumbnail": "CHAPTER 5\nUpdating Beliefs\nwith Evidence",
    },
    # 6 — Simpson's Paradox & Cost-Effectiveness
    {
        "filename": "06-simpsons-paradox-cost-effectiveness.docx",
        "title": "Chapter 6: Simpson's Paradox | When Combined Data Lies",
        "description": (
            "A treatment works in every subgroup — but fails overall. How is that possible? "
            "This is Simpson's Paradox, and it shows up more often than you think, "
            "especially when comparing health outcomes across countries.\n\n"
            "In Chapter 6, we learn to distinguish confounders from mediators, understand the "
            "Prosecutor's Fallacy, and apply causal reasoning to cross-country MNH comparisons. "
            "We also build R visualizations with ggplot2 to see the paradox in action.\n\n"
            "In this video, you will learn:\n"
            "- Simpson's Paradox: when aggregated data reverses the truth\n"
            "- Confounders vs. mediators — when to adjust and when not to\n"
            "- The Prosecutor's Fallacy — a dangerous reasoning mistake\n"
            "- Cross-country MNH comparisons using causal reasoning\n"
            "- Visualizing Simpson's Paradox in R with ggplot2\n\n"
            "TIMESTAMPS\n"
            "00:00 Introduction\n"
            "00:00 Simpson's Paradox explained\n"
            "00:00 A real-world health example\n"
            "00:00 Confounders vs. mediators\n"
            "00:00 The Prosecutor's Fallacy\n"
            "00:00 Cross-country comparisons\n"
            "00:00 R visualization with ggplot2\n"
            "00:00 Key takeaways\n\n"
        ),
        "tags": BASE_TAGS + [
            "Simpsons Paradox",
            "confounding variable",
            "mediator",
            "cost effectiveness",
            "ggplot2",
            "R programming",
            "data visualization",
            "Prosecutors Fallacy",
        ],
        "thumbnail": "CHAPTER 6\nSimpson's\nParadox",
    },
    # 7 — Decision Analysis & Value of Information
    {
        "filename": "07-decision-analysis-value-of-information.docx",
        "title": "Chapter 7: Decision Analysis | Watching vs. Doing in Health Programs",
        "description": (
            "Observing that something works is not the same as making it work. "
            "Chapter 7 introduces the crucial distinction between watching (inference) "
            "and doing (intervention) — and shows how graph surgery helps you predict "
            "what will actually happen when you act.\n\n"
            "We build decision maps, calculate expected values, and learn how to estimate "
            "the value of gathering more information before committing resources. "
            "Is it worth running a pilot first, or should you scale up now?\n\n"
            "In this video, you will learn:\n"
            "- Decision maps — laying out your options visually\n"
            "- Expected value — picking the option with the best average outcome\n"
            "- Watching vs. doing — observational vs. intervention effects\n"
            "- Graph surgery — cutting arrows to simulate interventions\n"
            "- Value of information — is more data worth the wait?\n\n"
            "TIMESTAMPS\n"
            "00:00 Introduction\n"
            "00:00 Decision maps\n"
            "00:00 Expected value calculations\n"
            "00:00 Watching vs. doing\n"
            "00:00 Graph surgery — cutting arrows\n"
            "00:00 Value of information\n"
            "00:00 MNH application\n"
            "00:00 Key takeaways\n\n"
        ),
        "tags": BASE_TAGS + [
            "decision analysis",
            "expected value",
            "value of information",
            "graph surgery",
            "do calculus",
            "intervention vs observation",
            "health program decisions",
        ],
        "thumbnail": "CHAPTER 7\nWatching\nvs. Doing",
    },
    # 8 — Multi-Country Resource Allocation
    {
        "filename": "08-multi-country-resource-allocation.docx",
        "title": "Chapter 8: Resource Allocation | Splitting a Budget Across Countries",
        "description": (
            "You have a limited budget and multiple countries that need support. "
            "Do you concentrate resources in one place, spread them evenly, or use a tiered "
            "approach? Chapter 8 tackles the resource allocation puzzle head-on.\n\n"
            "We compare three allocation strategies, introduce Monte Carlo simulation "
            '("replaying the future 1,000 times"), and explore sequential decision-making — '
            "adjusting your plan as you learn more. All examples use publicly available MNH data.\n\n"
            "In this video, you will learn:\n"
            "- Multi-country allocation as a puzzle with tradeoffs\n"
            "- Three strategies: Concentrate, Spread, and Tiered\n"
            "- Uncertainty and scenario analysis\n"
            "- Monte Carlo simulation — replaying the future 1,000 times\n"
            "- Sequential decision-making — adjusting as you learn\n\n"
            "TIMESTAMPS\n"
            "00:00 Introduction\n"
            "00:00 The allocation puzzle\n"
            "00:00 Strategy 1 — Concentrate\n"
            "00:00 Strategy 2 — Spread\n"
            "00:00 Strategy 3 — Tiered\n"
            "00:00 Uncertainty and scenarios\n"
            "00:00 Monte Carlo simulation\n"
            "00:00 Sequential decisions\n"
            "00:00 Key takeaways\n\n"
        ),
        "tags": BASE_TAGS + [
            "resource allocation",
            "Monte Carlo simulation",
            "budget allocation",
            "scenario analysis",
            "uncertainty",
            "sequential decision making",
            "health investment",
        ],
        "thumbnail": "CHAPTER 8\nResource\nAllocation",
    },
    # 9 — Game Theory & Strategic Interactions
    {
        "filename": "09-game-theory-strategic-interactions.docx",
        "title": "Chapter 9: Game Theory | When Other People React to Your Decisions",
        "description": (
            "Your decision is only half the story — what matters is how other people respond. "
            "Chapter 9 introduces game theory for health programs: payoff matrices, best responses, "
            "Nash Equilibrium, and the free-rider problem.\n\n"
            "We explore why some health partnerships fail even when everyone agrees on the goal, "
            "and learn about commitment devices and escape routes that can change the game.\n\n"
            "In this video, you will learn:\n"
            "- How other people's reactions change your outcomes\n"
            "- 2x2 payoff matrices and how to read them\n"
            "- Best responses and Nash Equilibrium\n"
            "- The free-rider problem in health partnerships\n"
            "- Commitment devices and escape routes\n\n"
            "TIMESTAMPS\n"
            "00:00 Introduction\n"
            "00:00 Why reactions matter\n"
            "00:00 Payoff matrices\n"
            "00:00 Best responses\n"
            "00:00 Nash Equilibrium\n"
            "00:00 The free-rider problem\n"
            "00:00 Commitment devices\n"
            "00:00 MNH partnership example\n"
            "00:00 Key takeaways\n\n"
        ),
        "tags": BASE_TAGS + [
            "game theory",
            "Nash equilibrium",
            "payoff matrix",
            "free rider problem",
            "strategic interaction",
            "commitment device",
            "health partnerships",
        ],
        "thumbnail": "CHAPTER 9\nGame Theory\nfor Health",
    },
    # 10 — Structure Learning from Observational Data
    {
        "filename": "10-structure-learning-from-data.docx",
        "title": "Chapter 10: Can Computers Discover Causes? | Structure Learning from Data",
        "description": (
            "Can an algorithm look at data and tell you what causes what? "
            "The answer is: partly. Chapter 10 explores what structure-learning algorithms "
            "can and cannot do — and why expert knowledge is still essential.\n\n"
            "We learn to combine algorithmic output with human judgment, build a data-model "
            "feedback loop, and implement structure learning in R using the bnlearn package's "
            "hill-climbing algorithm.\n\n"
            "In this video, you will learn:\n"
            "- Why data alone cannot determine the direction of arrows\n"
            "- What structure-learning algorithms can and cannot do\n"
            "- How to combine expert knowledge with algorithmic output\n"
            "- Data-model feedback loops — iterating between data and theory\n"
            "- R implementation with bnlearn's hill-climbing algorithm (hc)\n\n"
            "TIMESTAMPS\n"
            "00:00 Introduction\n"
            "00:00 Can data discover causes?\n"
            "00:00 What algorithms can do\n"
            "00:00 What algorithms cannot do\n"
            "00:00 Combining experts and algorithms\n"
            "00:00 The feedback loop\n"
            "00:00 R implementation with bnlearn::hc()\n"
            "00:00 Key takeaways\n\n"
        ),
        "tags": BASE_TAGS + [
            "structure learning",
            "observational data",
            "bnlearn",
            "hill climbing algorithm",
            "R programming",
            "causal discovery",
            "machine learning for health",
        ],
        "thumbnail": "CHAPTER 10\nCan Data\nDiscover Causes?",
    },
]

# ---------------------------------------------------------------------------
# Document builder
# ---------------------------------------------------------------------------


def add_heading(doc, text, level=1):
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        run.font.color.rgb = RGBColor(0, 0x5C, 0xB9)  # #005CB9
    return heading


def add_field(doc, label, value):
    """Add a bold label followed by the value on the same line."""
    p = doc.add_paragraph()
    run_label = p.add_run(f"{label}: ")
    run_label.bold = True
    run_label.font.size = Pt(11)
    run_value = p.add_run(value)
    run_value.font.size = Pt(11)
    return p


def add_multiline(doc, label, text):
    """Add a bold label then the text body below it."""
    p = doc.add_paragraph()
    run_label = p.add_run(f"{label}:")
    run_label.bold = True
    run_label.font.size = Pt(12)

    # Add description body
    body = doc.add_paragraph(text)
    body.style.font.size = Pt(11)
    return body


def build_doc(meta):
    doc = Document()

    # Title
    add_heading(doc, "YouTube Video Metadata", level=1)
    doc.add_paragraph("")  # spacer

    # Video title
    add_field(doc, "Video Title", meta["title"])
    doc.add_paragraph("")

    # Category
    add_field(doc, "Category", CATEGORY)
    doc.add_paragraph("")

    # Description
    full_desc = (
        meta["description"]
        + COURSE_FOOTER.format(course=COURSE_TITLE)
        + "\n\n"
        + HASHTAGS
    )
    add_multiline(doc, "Description", full_desc)
    doc.add_paragraph("")

    # Tags
    tag_str = ", ".join(meta["tags"])
    add_multiline(doc, "Tags", tag_str)
    doc.add_paragraph("")

    # Thumbnail text
    add_field(doc, "Suggested Thumbnail Text", meta["thumbnail"])

    return doc


# ---------------------------------------------------------------------------
# Main
# ---------------------------------------------------------------------------

def main():
    out_dir = os.path.dirname(os.path.abspath(__file__))
    for meta in VIDEOS:
        doc = build_doc(meta)
        path = os.path.join(out_dir, meta["filename"])
        doc.save(path)
        print(f"Created: {meta['filename']}")
    print(f"\nDone — {len(VIDEOS)} files created in {out_dir}")


if __name__ == "__main__":
    main()
