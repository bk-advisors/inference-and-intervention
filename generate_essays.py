"""
Generate Pre-Class Essays for "Inference and Intervention" Course
Output: pre-class-essays.docx
Style: Morgan Housel (conversational, story-driven, counterintuitive insights)
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
import os

# ── Essay Content ──────────────────────────────────────────────────────────────

essays = []

# ══════════════════════════════════════════════════════════════════════════════
# INTRODUCTION
# ══════════════════════════════════════════════════════════════════════════════
essays.append({
    "number": 0,
    "title": "The Gap Between Good Intentions and Good Outcomes",
    "subtitle": "A roadmap for the intellectual journey ahead",
    "body": r"""
In 2003, Ethiopia launched one of the most ambitious public health programs in African history. The government trained and deployed roughly 40,000 Health Extension Workers across the country. They built thousands of health posts. They hired women from local communities, gave them a year of training, and sent them into villages to provide basic care, health education, and referrals.

By almost any measure of effort, it was a staggering achievement. But when researchers looked at the results district by district, they found something unsettling. Some of the regions with the highest density of health workers also had the highest rates of child mortality. More workers. Worse outcomes. How could that be?

The answer was not that the health workers were harmful. The answer was that the government had sent more workers to the places that needed them most -- the remote, underserved regions where mortality was already sky-high and where roads, clean water, and functioning clinics were scarce. The workers were swimming against a current that was stronger than their training.

This is not a story about failure. Ethiopia's Health Extension Program has saved hundreds of thousands of lives over two decades. It is one of the great public health achievements of the twenty-first century.

But it is a story about how looking at data without thinking carefully about cause and effect can lead you to exactly the wrong conclusion. If you saw the correlation between worker density and mortality, and you took it at face value, you might recommend pulling workers out of the hardest-hit regions. You would be making the problem worse while believing you were making it better.

This gap -- between what the data seems to say and what is actually true -- is where this course lives.

We are going to spend ten chapters learning a set of tools for thinking about cause and effect. These tools come from a field called causal inference, and they are some of the most important intellectual inventions of the past fifty years. But they are not widely taught, and they are almost never taught to the people who need them most: the decision-makers, program managers, and policy advisors who allocate resources, design interventions, and evaluate results in global health.

Here is a preview of the journey.

We start with the most basic distinction in all of causal reasoning: the difference between a correlation and a cause. Chapter 1 uses Ethiopia's health worker data to show how confounding, reverse causation, and selection bias can make innocent patterns look guilty and guilty patterns look innocent. You will learn why you should never trust a number until you have drawn a picture of how you think the world works.

Chapter 2 teaches you how to draw that picture. Using Rwanda's community health worker program as a case study, you will learn the grammar of causal diagrams -- the three fundamental structures that govern how information flows between variables. You will discover that the simple act of putting your beliefs on paper is one of the most productive things you can do before touching any data.

Chapter 3 takes you into the field. Using Kenya's 47 counties as a laboratory, you will learn how to build a causal model through interviews -- starting with a simple sketch and refining it as each conversation reveals what the last one missed.

In Chapters 4 through 6, we add numbers to the picture. Chapter 4 introduces probability as a language for expressing what you do not know. Chapter 5 shows you how to run the model backwards -- how observing an outcome changes your beliefs about its causes, including the strange phenomenon of explaining away. Chapter 6 reveals one of the most dangerous traps in data analysis: Simpson's Paradox, where aggregated data across countries can recommend the exact opposite of what disaggregated data shows.

Chapters 7 through 9 are where the course gets practical. Chapter 7 introduces the formal distinction between watching something happen and making it happen -- between observing that Kenyan counties with CPAP machines have better outcomes and actually giving a CPAP machine to a county that lacks one. Chapter 8 applies portfolio theory to resource allocation across four countries. Chapter 9 adds the hardest layer: other people have strategies too. When you invest in a country's health system, the government responds. Sometimes it matches your investment. Sometimes it pockets the savings.

Chapter 10 closes the loop by asking whether data alone can discover causes. The answer is a qualified yes -- algorithms can find some causal structures, but they cannot find all of them. You will learn why expert knowledge and data are not substitutes. They are complements.

That is the journey. From a misleading correlation in Ethiopia to a formal, testable, updatable framework for thinking about cause and effect in the world's most complex systems.

Why does it matter?

Because every year, billions of dollars flow into global health programs based on analyses that confuse correlation with causation. Programs get scaled based on patterns that reverse when you look more carefully. Resources get concentrated in places that look effective on paper but are simply easier to reach. And governments respond to donor investments in ways that no regression table will ever capture.

These are not rare mistakes made by careless people. They are systematic errors built into the standard analytical toolkit. And they can only be corrected by learning to think causally.

The good news is that causal thinking is a skill, not a talent. You can learn it. And once you do, you will never look at a number the same way again.
""",
})

# ══════════════════════════════════════════════════════════════════════════════
# CHAPTER 1
# ══════════════════════════════════════════════════════════════════════════════
essays.append({
    "number": 1,
    "title": "The Most Dangerous Number Is a Correlation",
    "subtitle": "Why trusting data without a causal model can lead you to recommend the worst possible option",
    "body": r"""
Let us go back to Ethiopia's Health Extension Program and look at the numbers more carefully.

In the mid-2000s, the Ethiopian government rolled out tens of thousands of community-based health workers. The deployment was not random. Remote, underserved areas received more workers per capita. Urban areas with existing clinics received fewer. The idea was simple: send help where it is needed most.

Now imagine you are an analyst sitting in an office, looking at a spreadsheet. You have two columns: health worker density by district, and under-five mortality by district. You plot them. You run a correlation. And you find that districts with more health workers per capita tend to have higher mortality rates.

If you stop here, you have a dangerous number. You have a fact -- a real, reproducible statistical fact -- that points in exactly the wrong direction. It suggests that health workers are associated with death, not life.

Of course, the health workers did not cause the higher mortality. The relationship ran the other way around. High mortality was the reason those districts received more workers. The government's allocation rule created a pattern in the data that looked like harm but was actually help arriving at the scene of an emergency.

This is confounding. A third variable -- the underlying difficulty of the district, its remoteness, its poverty, its lack of infrastructure -- drove both the number of health workers deployed and the mortality rate observed. The health workers and the mortality were correlated, but neither caused the other. They were both consequences of the same root cause.

Here is the uncomfortable truth about confounding: it does not look any different from a real causal effect in the data. The correlation between health workers and mortality is just as strong, just as statistically significant, and just as reproducible as a genuine causal relationship would be. There is nothing in the numbers themselves that tells you it is spurious. The data does not come with a warning label.

This is why we need causal models -- pictures drawn before we look at the data that represent our beliefs about how the world works. A causal model for this situation would show an arrow from "district difficulty" to "health worker deployment" and another arrow from "district difficulty" to "mortality." It would show no direct arrow from "health worker deployment" to "mortality" -- or if it did, the arrow would point in the direction of reducing mortality, not increasing it. The model would immediately reveal that the observed correlation is driven by a backdoor path through district difficulty, not by a direct causal link.

Without that model, the correlation is a trap.

Confounding is the first of three traps that catch analysts who skip the causal model. The second is reverse causation.

Consider a pattern you can find in data from many African countries: districts with more doctors tend to have better health outcomes. A natural interpretation is that doctors improve health. And they do. But there is another story running simultaneously. Districts with better health outcomes -- wealthier districts, more urbanized districts, districts with better infrastructure -- also attract more doctors. Doctors prefer to live where the schools are good, the roads are paved, and the electricity is reliable.

So which is it? Do doctors cause better health, or does better health attract doctors? The answer, almost certainly, is both. But the relative strength of these two effects matters enormously for policy. If the main story is that doctors cause health improvements, you should send doctors to underserved areas. If the main story is that good conditions attract doctors, you need to fix roads and schools before the doctors will come. The arrow has to point one way or the other, and the policy recommendation depends on which way it points.

In cross-sectional data -- a single snapshot of all districts at one point in time -- you cannot tell these two stories apart. They produce identical correlations. Only a causal model, combined with the right kind of data or the right kind of study design, can distinguish them.

The third trap is selection bias, and it is the sneakiest.

Imagine a donor funds a competitive grant program for health facilities in Tanzania. Facilities apply, reviewers score them, and the top scorers receive funding. Two years later, someone evaluates the grant recipients and finds that facilities with more experienced staff had better outcomes. The conclusion: experience matters. Hire experienced people.

But wait. The grant process selected for facilities that already had a combination of strong leadership, experienced staff, and credible plans. Among this selected group, experience and outcomes are correlated partly because the selection filter linked them. If you looked at all facilities in the country -- not just grant recipients -- the relationship between experience and outcomes might be much weaker, or structured differently.

By conditioning on grant receipt -- by only looking at the winners -- you created a distorted picture of how experience relates to outcomes in general. You selected on a collider, and the result was a phantom pattern that does not exist in the broader population.

These three traps share something important: they all produce real statistical patterns that feel meaningful but are not causal. The correlations are not noise. They are not errors. They are genuine features of the data. They just do not mean what you think they mean.

The defense against all three is the same. Before you look at any data, draw a picture. Put boxes for the variables you care about. Put arrows for the causal relationships you believe exist. Then ask: given this picture, what patterns should I expect to see in the data? And what patterns would be misleading?

The picture is your causal model. It is the most important analytical tool you will ever learn to use. And it costs nothing to draw except honesty -- the willingness to say, "This is what I believe, and I might be wrong."

The most dangerous number is not a wrong number. It is a right number with the wrong story attached to it. And the only way to check the story is to draw the picture first.
""",
})

# ══════════════════════════════════════════════════════════════════════════════
# CHAPTER 2
# ══════════════════════════════════════════════════════════════════════════════
essays.append({
    "number": 2,
    "title": "Drawing What You Believe",
    "subtitle": "Why writing down your causal beliefs is the most productive argument you will ever have",
    "body": r"""
In 2005, Rwanda rebuilt its health system almost from scratch. The country had emerged from genocide a decade earlier with barely a functioning clinic. By the mid-2000s, the government had designed one of the most structured community health worker programs in Africa. Every village got three workers. Every worker got defined tasks. Every task got tracked.

But here is what is often missed about Rwanda's success: different experts looking at the same program had completely different theories about why it worked.

The public health specialist believed the key was coverage. Three workers per village meant every household was within walking distance of basic care. The economist believed the key was incentives. Rwanda paid its community health workers performance-based bonuses tied to measurable outcomes. The political scientist believed the key was governance. Rwanda's top-down administrative structure meant that policies set in Kigali were actually implemented in villages, which is far from guaranteed in most countries.

Each expert had a mental model. Each model led to a different recommendation for other countries trying to replicate Rwanda's results. And each expert could point to data supporting their view.

The problem was not that they disagreed. The problem was that they were arguing about conclusions while their disagreements were really about assumptions. They were debating what to do without ever surfacing why they believed different things would work.

A causal diagram fixes this. Not by telling you who is right, but by forcing everyone to draw their map -- and then comparing the maps side by side.

A causal model is a diagram with two ingredients. Nodes represent things you care about: coverage, incentives, governance, training quality, mortality. Arrows represent your beliefs about what influences what. An arrow from "incentives" to "worker effort" means you believe that changing incentives would change effort. Not a correlation. A genuine causal claim. If you could reach into the system and adjust incentives, you believe effort would shift.

The act of drawing this diagram is deceptively powerful. Here is why.

When your beliefs live inside your head, they do not have to be consistent. You can believe that "training improves quality" and also that "quality is mainly determined by equipment" without ever noticing the tension. You never hold both beliefs up to the light at the same time. They coexist peacefully in the dark.

The moment you draw an arrow from "training" to "quality" and try to decide whether you also need an arrow from "equipment" to "quality" -- and if so, how strong each arrow is relative to the other -- the tension becomes visible. The diagram demands specificity that intuition avoids.

There are four types of nodes, and they matter more than you might expect.

Probabilistic nodes represent things that are uncertain. Will the government sustain funding? Will community health workers stay in their posts? You cannot control these. You can influence them, maybe, but you cannot set them with certainty. These are the unknowns that make decision-making hard.

Objective nodes represent what you are trying to achieve. Lives saved. Cost-effectiveness. These are your destinations. Everything else in the model exists in relation to them.

Decision nodes represent things you can control. Fund training or fund equipment? Invest in one district or spread across ten? These are your levers.

Function nodes represent quantities that are fully determined by their inputs. If you know the number of workers and the cost per worker, you know the total cost. No uncertainty, just arithmetic. These seem boring, but they keep the model honest by separating genuine uncertainty from simple calculation.

Now, the arrows create structures, and these structures have profound consequences for what you can and cannot learn from data. There are exactly three fundamental structures. If you understand these three, you understand the grammar of causation.

The first is the chain: A causes B, and B causes C. Rwanda invests in worker training. Training improves care quality. Quality reduces mortality. In this structure, training and mortality are correlated -- they move together in data -- but only because quality carries the signal. If you could somehow hold quality constant, training and mortality would become independent. The middleman is doing all the work.

The second is the fork: B causes both A and C. Strong governance leads to both good supply chains and effective supervision. In data, supply chain performance and supervision quality will be correlated -- not because one causes the other, but because they share a common parent. Hold governance constant, and the correlation vanishes. They were both reflections of the same underlying cause.

The third structure is the one that breaks intuition: the collider. A and B both cause C, but A and B are otherwise unrelated. Suppose staff motivation and equipment quality both affect patient outcomes. Motivation and equipment are independent in general. But if you only look at facilities with good outcomes -- you condition on C -- suddenly motivation and equipment become negatively correlated. Among the successes, the ones that succeeded despite bad equipment must have had great staff, and vice versa. Conditioning on the outcome creates a phantom link between its causes.

This collider effect is behind a stunning number of analytical mistakes. It is why studying only successful programs gives you misleading predictors of success. It is why looking only at hospitals that received grants can make unrelated factors appear connected. Conditioning on the wrong variable does not just fail to help. It actively misleads.

These three structures -- chains, forks, and colliders -- are the atoms of causal reasoning. Every model, no matter how many nodes and arrows it has, is built from combinations of these three patterns.

And this brings us back to that room in Kigali where three experts cannot agree about Rwanda. When you ask each one to draw a causal model, the disagreements become productive. Instead of arguing about what to fund, they argue about arrows. "Do you really think governance directly affects mortality, or does it work through supervision and supply chains?" "Is worker motivation caused by incentives, or by community respect, or by both?"

These are questions you can investigate. They are a thousand times more useful than arguing about which program to fund.

Draw what you believe. Then be prepared to discover that your colleagues believe something different. That discovery is not a problem. It is the beginning of real analysis.
""",
})

# ══════════════════════════════════════════════════════════════════════════════
# CHAPTER 3
# ══════════════════════════════════════════════════════════════════════════════
essays.append({
    "number": 3,
    "title": "The Art of Asking Better Questions",
    "subtitle": "How each conversation reveals what the previous model could not see",
    "body": r"""
Kenya has 47 counties. Each one runs its own health system. Since devolution in 2013, counties have had the authority -- and the responsibility -- to fund clinics, hire staff, buy supplies, and set priorities. Some counties have thrived. Others have struggled. And the variation is enormous.

If you wanted to understand why neonatal outcomes differ so dramatically across Kenyan counties, you could start with the data. Kenya's District Health Information System has millions of records. You could run regressions for weeks.

But here is what the data will not tell you: why. It will tell you that some counties have more skilled birth attendants per capita. It will not tell you why those counties managed to recruit and retain them. It will tell you that some counties have lower stockout rates. It will not tell you whether the supply chain works because of a competent county pharmacist or because of a functioning logistics system or because of political pressure from the governor.

The data shows patterns. The interviews reveal mechanisms. And mechanisms are what you need to build a causal model.

Imagine a team setting out to build a causal model of neonatal mortality across Kenyan counties. They start where anyone would start: with the obvious variables. Skilled birth attendance. Facility readiness. Referral systems. They draw four nodes on a whiteboard and connect them with arrows. It is clean. It is logical. It fits on a napkin.

Then they sit down with the Director of Health in Makueni County.

Within twenty minutes, the model is in trouble. The Director mentions something the team had not considered: county budget execution rates. "We receive allocations from the national government," she explains, "but the money often arrives late. Sometimes very late. And you cannot hire staff or buy supplies with money you do not have yet." She adds: "The counties that perform well are often the ones with governors who prioritize health. The governor's priorities determine how fast the money moves."

The four-node model just became a seven-node model. Budget execution. Political prioritization. Timeliness of fund disbursement. Three new variables that were invisible from the national data.

For the second round, the team visits three facilities in different counties -- one high-performing, one average, one struggling.

The high-performing facility is in Nyeri County. The clinical officer talks about something the team had not expected: mentorship networks. "We have a WhatsApp group with other clinical officers across the county. When I have a complicated case at night, I send a photo and get advice within minutes. That network has saved more lives than any piece of equipment in this building."

Peer networks. Informal knowledge sharing. These do not appear in any facility assessment checklist. But they might be among the most important factors in care quality.

The struggling facility is in Turkana County. The nurse in charge paints a bleak picture. "My biggest challenge is not equipment or training. It is that mothers do not come. They deliver at home with a traditional birth attendant because the nearest facility is four hours away on foot, and there is no transport. Even when something goes wrong, by the time they reach us, it is often too late."

Geographic access. Transport infrastructure. Demand-side barriers. The model grows again. And the arrows are different from what the team expected. The problem is not inside the facility. It is between the community and the facility door.

The average facility, in Kisumu County, reveals yet another layer. The facility manager says: "We have good staff. We have decent equipment. Our problem is data. We report to the county, the county reports to the national level, but nothing comes back. We do not know how we compare to other facilities. We do not know if our referral rates are normal or alarming. We are flying blind."

Feedback loops. Information systems. The model now includes the idea that data is not just something analysts use to evaluate programs. It is something frontline workers need to improve their own performance. And its absence might be a cause of stagnation.

Round three takes the team to the community. They interview community health volunteers -- the people who visit pregnant women in their homes.

A volunteer in Kilifi County explains: "I am supposed to track every pregnant woman in my area and make sure she attends antenatal care at least four times. But I cover 500 households on foot. I have no phone. I have no register. I remember who I visited last week and try to rotate, but I know I am missing people."

Another volunteer, in Bungoma County, describes a different challenge: "The women trust me, but their husbands do not want to spend money on transport to the facility. He says, 'My mother delivered at home and she was fine.' I am not fighting ignorance. I am fighting a household decision-making structure where the woman's health is not the priority."

The model now has fifteen or twenty nodes. It includes male involvement in health decisions. Community health volunteer workload. Data feedback loops. Governor priorities. Transport infrastructure. Peer mentorship networks. It is messy. It is complicated. And it is incomparably more useful than the four-node sketch the team started with.

This iterative process -- build a model, test it against reality, revise, repeat -- is the heart of causal analysis. Each interview does two things. First, it adds new nodes and arrows that the team could not have imagined from their office. Second, it restructures existing relationships. An arrow the team drew as direct turns out to be mediated. A relationship they assumed was universal turns out to depend on context.

There is a temptation, especially among quantitative analysts, to skip the interviews and go straight to the data. We have Kenya's DHIS2. We have Demographic and Health Surveys. We have facility assessments. Why spend weeks talking to people?

Because data can tell you what is correlated with what. It cannot tell you why. It cannot tell you that money arrives late because of county politics. It cannot tell you that mothers stay home because of their husbands, not because of ignorance. It cannot tell you that a WhatsApp group is saving more lives than a ventilator.

And it cannot tell you what is missing from your model. You cannot include a variable in your regression if you do not know it exists. And you cannot know it exists if you have not talked to the people who live inside the system.

Your first model will be wrong. Your second will be less wrong. Your third will be less wrong still. The goal is not perfection. It is progressive honesty about the system you are trying to change.
""",
})

# ══════════════════════════════════════════════════════════════════════════════
# CHAPTER 4
# ══════════════════════════════════════════════════════════════════════════════
essays.append({
    "number": 4,
    "title": "Putting Numbers on Uncertainty",
    "subtitle": 'Why "I do not know" is the beginning of analysis, not the end',
    "body": r"""
Tanzania has roughly 7,000 health facilities. According to national surveys, only about half of them have the essential medicines needed to treat the most common causes of death in newborns. When you read that statistic, what does it actually mean?

It could mean that roughly half of all facilities are fully stocked and half are completely empty. It could mean that nearly every facility has some medicines but none has all of them. It could mean that facilities in urban areas are well-supplied while rural facilities have almost nothing. The single number -- "about half" -- hides an enormous amount of variation, and the variation matters for what you should do next.

If the problem is that a small number of facilities have nothing, you send supplies. If the problem is that almost every facility is missing one or two critical items, you fix the distribution system. If the problem is geographic, you fix the supply chain to rural areas. Same statistic. Three different diagnoses. Three different interventions.

This is why you need to put numbers on uncertainty -- not because the numbers will be perfectly right, but because the act of specifying what you know and what you do not know forces a precision that vague descriptions avoid.

In the previous chapters, we built qualitative causal models -- diagrams with arrows showing what causes what. Those models captured the structure of our beliefs. Now we need to capture their strength. You believe training affects quality. But how much? You believe supply chains affect medicine availability. But with what reliability?

A probability is a number between 0 and 1 that represents your degree of belief that something is true. If you say the probability of a facility having essential medicines is 0.5, you are saying you think it is a coin flip. If you say 0.8, you are fairly confident. If you say 0.2, you are fairly confident it does not.

These numbers are not just mathematical abstractions. They are commitments. When you write down a probability, you are staking a claim that can be checked against evidence, debated with colleagues, and updated when new information arrives.

Conditional probability is where things get powerful. The question is not just "What is the probability of good care quality?" It is "What is the probability of good care quality, given that this facility has well-trained staff?" And then: "What is the probability of good care quality, given that the staff are well-trained but the essential medicines are missing?"

These conditional probabilities are organized in what we call Conditional Probability Tables, or CPTs. For each node in your causal model, you write down the probability of each state of that node for every combination of its parent states.

Suppose care quality depends on two things: staff training (high or low) and medicine availability (adequate or inadequate). You need four numbers. What is the probability of high quality when training is high and medicines are adequate? Maybe you say 85 percent. When training is high but medicines are inadequate? Maybe 50 percent. When training is low but medicines are adequate? Maybe 40 percent. When both are low? Maybe 10 percent.

These numbers force you to make real commitments about how you think the world works. They reveal beliefs you did not know you had. Do you think training matters more than medicines, or the other way around? Do they compensate for each other, or do you need both? The CPT forces you to answer.

Now, Bayes' Rule. If there is one idea that changes how you think about evidence, this is it.

Imagine you are a regional health officer in Tanzania. You receive a report that a newborn died at a facility in your district. You want to know: what went wrong? Your causal model has several possible explanations -- the medicine was unavailable, the staff lacked training, the mother arrived too late, or the case was simply too severe for any facility to manage.

You have some prior beliefs about how common each of these problems is. You also know something about how likely each problem is to produce the outcome you observed. Bayes' Rule tells you how to combine these two pieces of information. It says: the probability that a cause is responsible, given the outcome you observed, is proportional to how common that cause is (the prior) multiplied by how well it explains the outcome (the likelihood).

A rare cause that perfectly explains the evidence can outweigh a common cause that only weakly explains it. This is the logic of diagnosis. It is how doctors think -- or should think -- when they see symptoms. And it is how you should think when you see patterns in health data.

The Causal Markov Condition ties this all together. It says: given its direct causes, a variable is independent of everything else that is not its descendant. In simple terms, once you know what directly causes something, knowing about its grandparents or distant relatives tells you nothing extra.

This sounds technical, but it is actually a gift. It means you can build a complex model with dozens of nodes and still compute probabilities efficiently. Instead of needing one massive table covering every possible combination of every variable, you only need small local tables -- one per node -- that describe how each variable relates to its direct parents. The Causal Markov Condition guarantees that these local tables, combined with the graph structure, capture everything you need.

This is a Bayesian network: a causal diagram plus a set of CPTs. The diagram tells the qualitative story -- what causes what. The CPTs tell the quantitative story -- how strongly, and with what probability.

Let us return to Tanzania's essential medicines problem. In a Bayesian network, you might have nodes for national procurement, regional distribution, facility storage, and medicine availability. Each node has a CPT that specifies how its state depends on its parents. National procurement affects regional distribution. Regional distribution affects facility storage. Facility storage affects whether the medicine is actually available when a patient needs it.

When you hear that a facility in Mtwara region lacked essential medicines, you can run the model and ask: given this observation, what is the most likely bottleneck? Is it national procurement (maybe the medicine was never purchased)? Regional distribution (maybe it was purchased but never shipped)? Or facility storage (maybe it arrived but was stored improperly or expired)?

The answer depends on your CPTs -- on the numbers you put in. And the numbers might be uncertain. That is fine. The point is not to get the numbers perfect. The point is to make your reasoning transparent, updatable, and internally consistent.

Building a Bayesian network is an exercise in structured honesty. You do not need to know everything. You just need to say, for each variable, what you believe about its relationship with its direct causes. "I am not sure" is a perfectly valid starting point. You express that uncertainty as a probability near 0.5, and you update it as evidence comes in.

The payoff is enormous. Once the network is built, it can do things your intuition cannot. It can propagate evidence through chains of variables. It can compute the value of gathering additional information before making a decision. It can handle the counterintuitive dynamics we will explore in the next chapter.

Uncertainty is not the enemy of good decision-making. It is the raw material. The trick is to measure it.
""",
})

# ══════════════════════════════════════════════════════════════════════════════
# CHAPTER 5
# ══════════════════════════════════════════════════════════════════════════════
essays.append({
    "number": 5,
    "title": "Running the Model Backwards",
    "subtitle": "How observing an outcome changes everything you think about its causes",
    "body": r"""
There is a region in Ethiopia -- let us call it the kind of region that should be succeeding. It has received above-average government funding. It has a reasonable number of health workers. It has functional supply chains. On paper, it looks like a region where neonatal outcomes should be improving.

But the outcomes are not improving. The mortality rate has stalled. Neighboring regions with fewer resources are doing better.

If you are a program manager looking at this situation, your natural instinct is to ask: what is going wrong? And your natural instinct is to look for the most obvious explanation -- maybe the funding was mismanaged, maybe the workers are poorly trained, maybe the data is wrong.

But there is a more disciplined way to approach this question. Instead of guessing, you can run your causal model backwards. You can enter what you observe -- stalled outcomes despite adequate resources -- and let the model tell you what that evidence implies about every variable in the system.

This is called inference, and it produces results that are often surprising.

Here is the core idea. Your Bayesian network has nodes for government funding, health worker deployment, supply chain reliability, care quality, community engagement, and neonatal outcomes. Each node has a conditional probability table expressing your beliefs about how it depends on its parents. Normally, you run the model forward: you start with inputs and predict outputs. But the mathematics of Bayes' Rule lets you run it in reverse: you start with an observed output and update your beliefs about the inputs.

When you enter the evidence -- "outcomes are poor despite apparently adequate funding and staffing" -- the model updates your beliefs about every node simultaneously. And one of the most important things it does is called explaining away.

Explaining away works like this. Suppose neonatal outcomes in this region depend on two things: care quality at facilities and community engagement with the health system. Before you know anything about the outcome, these two causes are independent in your model. Knowing about care quality tells you nothing about community engagement, and vice versa.

Now you observe that outcomes are poor. This is bad news, and it makes you update your beliefs about both potential causes downward. But then you get more information: a facility assessment shows that care quality is actually decent. The staff are trained. The equipment works. The protocols are followed.

Here is where things get counterintuitive. The moment you learn that care quality is adequate, your belief that community engagement is the problem goes sharply up. Not because you have any direct evidence about community engagement. But because the poor outcome has to be explained by something, and you have just ruled out the most obvious candidate. The remaining explanation absorbs the burden.

This is explaining away. Two independent causes become linked once you observe their common effect. Confirming one cause reduces your belief in the other. Ruling out one cause increases your belief in the remaining ones.

If this feels strange, think of it like a detective story. A crime has been committed. You have two suspects, initially unrelated. Then you learn that one suspect has an alibi. Suddenly, the other suspect looks much more guilty -- not because you have new evidence against them, but because the crime still needs an explanation, and the list of candidates just got shorter.

In our Ethiopian region, the model might point strongly toward community engagement as the bottleneck. Perhaps mothers are not delivering at facilities. Perhaps community health workers are present but not actively reaching out to pregnant women. Perhaps cultural factors are keeping families away from the health system. None of this was directly observed. The model inferred it by ruling out the alternatives.

This kind of reasoning -- combining multiple pieces of evidence, some confirming and some contradictory, through a causal structure -- is something humans do poorly in their heads. We tend to anchor on the most vivid or most recent piece of evidence and ignore the rest. We tend to confirm our first hypothesis rather than systematically evaluating alternatives. The Bayesian network does the math correctly and consistently.

Let us look at another example. Suppose you are comparing two districts that both have poor neonatal outcomes.

In the first district, you find that government funding is low, health worker deployment is thin, and supply chains are unreliable. The model is not surprised. The poor outcomes are explained by a cascade of resource deficits. There is no mystery. The diagnosis is straightforward: more resources.

In the second district, you find that government funding is reasonable, health workers are deployed, but outcomes are still poor. The model now has to work harder. It updates your beliefs about the things you have not observed. Maybe care quality is poor despite adequate inputs -- perhaps the workers are present but poorly supervised. Maybe the supply chain is technically functional but delivers the wrong supplies. Maybe demand-side barriers are preventing mothers from using the services that exist.

The same model, applied to different evidence, produces different diagnoses. This is the power of running the model backwards. It does not apply a universal formula. It adapts to the local context, processing evidence through a causal structure to identify the most likely bottleneck in each specific situation.

There is a broader lesson here about how we process evidence in everyday life. We are all running models backwards all the time, usually without realizing it. When a student does well on an exam, we update our beliefs about their intelligence and their effort. If we then learn they studied all night, our estimate of their natural ability goes down -- explaining away. The success was accounted for by effort, reducing the need to invoke talent.

When a health program succeeds in a difficult environment, we give credit to program design. But if we then learn that the government was exceptionally supportive, the credit we give to program design decreases. The success was explained by context, not just by the program.

The problem is that we do this intuitively and inconsistently. We explain away too aggressively when we like one hypothesis and not aggressively enough when we do not. We anchor on emotionally satisfying explanations. We forget that observing an outcome fundamentally changes the relationship between its causes.

A formal model does not have these biases. It updates beliefs exactly as the laws of probability require. It does not care which explanation is politically convenient or emotionally satisfying. It follows the arrows and the numbers.

This does not mean the model is always right. It is only as good as its structure and its probabilities. But it provides a disciplined, transparent starting point that human judgment can then refine.

Run the model backwards. You might be surprised by what you learn about the causes you have not observed yet.
""",
})

# ══════════════════════════════════════════════════════════════════════════════
# CHAPTER 6
# ══════════════════════════════════════════════════════════════════════════════
essays.append({
    "number": 6,
    "title": "When the Average Lies",
    "subtitle": "Why aggregated data can recommend the exact wrong thing -- and how to catch it",
    "body": r"""
Suppose you are comparing community health worker programs across four East African countries -- Kenya, Rwanda, Ethiopia, and Tanzania. You have outcome data -- let us say neonatal mortality reduction -- for each country's program. The aggregate numbers look clear. Kenya achieved a 30 percent reduction. Rwanda achieved 20 percent. Ethiopia achieved 15 percent. Tanzania achieved 10 percent. The ranking seems obvious. Learn from Kenya.

But now you break the data down by facility type -- hospitals versus health centers versus dispensaries. And the picture flips.

Among hospitals, Tanzania achieved the best outcomes. Among health centers, Tanzania is also in the lead. Among dispensaries, Tanzania beats everyone. Tanzania is the best performer at every level of facility -- and yet, in the aggregate, it ranks last.

How is this possible?

The answer is Simpson's Paradox. Tanzania's program operates overwhelmingly in dispensaries, which are the hardest settings -- rural, understaffed, under-equipped. Its aggregate number is dragged down by the difficulty of its patient mix. Kenya's program operates mostly in hospitals, where outcomes are better regardless of the program. Kenya's high aggregate number is a reflection of where it works, not how well it works.

The aggregate data did not just fail to identify the most effective program. It recommended the worst one. If you scaled Kenya's approach to dispensary settings, you would get worse results than Tanzania's approach. The average lied.

This is not a hypothetical. Simpson's Paradox appears whenever you aggregate data across groups that differ in both their treatment assignment and their baseline risk. In global health, this is almost always the case. Countries have different disease burdens, different baseline capacities, different levels of urbanization. Programs that operate in harder settings will look worse in aggregate even if they are more effective within every setting.

Let us trace the logic to understand why this happens.

A confounder is a variable that influences both the treatment and the outcome. In our example, the type of health facility is the confounder. It influences which country's approach is used (because different countries operate in different settings) and it influences outcomes (because hospitals produce better outcomes than dispensaries regardless of the program). When you aggregate across facility types, the confounder is hiding in the mix, silently biasing the comparison.

When you disaggregate by facility type, you are holding the confounder constant. You compare hospitals to hospitals, dispensaries to dispensaries. Now the comparison is fair, and the true ranking emerges.

But here is the harder question: should you always disaggregate? Is the within-group analysis always right and the aggregate always wrong?

No. And this is where the causal model becomes essential.

If the variable you are disaggregating on is a confounder -- a common cause of both the treatment and the outcome -- then you should disaggregate. Failing to do so produces a biased estimate.

But if the variable is a mediator -- something that sits on the causal pathway between the treatment and the outcome -- then disaggregating is wrong. You would be removing part of the treatment's effect.

Consider this: a training program improves health worker knowledge, which improves care quality, which reduces mortality. Knowledge is a mediator. If you compare trained and untrained workers while holding knowledge constant, you will find that training has no effect -- because you have blocked the very pathway through which training works. You have controlled for the mechanism and then wondered why the mechanism disappeared.

The difference between a confounder and a mediator is not visible in the data. They can produce identical statistical patterns. The difference is in the causal model -- in the direction of the arrows. A confounder points into both the treatment and the outcome. A mediator points from the treatment through to the outcome. Only the causal story tells you which is which.

This is why Simpson's Paradox is not really a statistical problem. It is a causal problem. The data alone cannot tell you whether to aggregate or disaggregate. You need a model.

There is a related trap that shows up in diagnostic testing, sometimes called the Prosecutor's Fallacy.

Suppose you have a screening test for preeclampsia. It catches 95 percent of true cases and correctly clears 90 percent of non-cases. An administrator sees these numbers and concludes: this test is excellent. Mandate it everywhere.

But what is the probability that a woman actually has preeclampsia if she tests positive?

It depends on the base rate. If preeclampsia affects 5 percent of the population, then among 1,000 women, 50 have the condition and 950 do not. The test catches 47 or 48 of the true cases. But it also falsely flags 95 of the healthy women. So out of roughly 143 positive results, only about 48 are real. A "95 percent accurate" test is actually wrong about two-thirds of the time when the condition is rare.

This is not a flaw in the test. It is a consequence of base rates. And confusing test accuracy with the probability of disease given a positive test is the Prosecutor's Fallacy. It has sent innocent people to prison. In global health, it leads to wasted resources on false positives and misplaced confidence in screening programs.

The cure is Bayes' Rule. Always combine test performance with the prior probability of the condition.

Both Simpson's Paradox and the Prosecutor's Fallacy emerge when you ignore the structure of the problem and focus only on surface-level numbers. Aggregate success rates and test accuracy percentages are not wrong numbers. They are incomplete numbers. They answer a different question than the one you are asking, and the distance between those questions can be measured in lives.

When you compute cost-effectiveness across countries, you are aggregating across wildly different contexts. The aggregate number might say that investing in Country X is twice as cost-effective as Country Y. But disaggregated by facility type, the ranking might flip. The causal model tells you which analysis is correct.

The average is the most natural statistic in the world. It is also, in the wrong context, the most dangerous one. Trust it only after you have checked what it is hiding.
""",
})

# ══════════════════════════════════════════════════════════════════════════════
# CHAPTER 7
# ══════════════════════════════════════════════════════════════════════════════
essays.append({
    "number": 7,
    "title": "The Difference Between Watching and Doing",
    "subtitle": "Why giving a county a CPAP machine is not the same as noticing it already has one",
    "body": r"""
Here is a sentence that sounds like a tautology but is actually one of the deepest ideas in all of causal reasoning:

Seeing that something is true is not the same as making it true.

Let me show you why this matters with a real pattern from Kenya.

Kenya's 47 counties vary enormously in their health infrastructure. Some have neonatal intensive care units with functioning CPAP machines -- devices that help newborns breathe. Others have nothing. If you look at the data, counties with CPAP machines have significantly better neonatal outcomes than counties without them. The correlation is strong. It survives adjustment for obvious factors like county population and urban-rural mix.

A straightforward conclusion: CPAP machines save lives. Buy more CPAPs.

But wait. Why do some counties have CPAPs and others do not? The counties with CPAPs tend to be wealthier counties. They have stronger tax bases, more political visibility, better-managed hospitals, and more skilled staff. The CPAP machine is not sitting in a vacuum. It is part of a package -- a well-functioning system that includes trained clinicians, reliable electricity, oxygen supplies, and maintenance capacity.

When you observe that a county has a CPAP, you are seeing the product of a system. That system placed the CPAP there, and the same factors that drove the placement also drive outcomes. The observation is contaminated by the process that generated it.

When you intervene to give a CPAP to a county that does not have one -- a rural county in arid northern Kenya where the electricity is unreliable, the oxygen supply is sporadic, and the nearest trained neonatologist is 300 kilometers away -- you are doing something fundamentally different. You are bypassing the system. You are putting a CPAP in a setting that the system would not have selected. The county gets the machine, but it does not get the constellation of factors that typically accompanies it.

This distinction -- between observing a variable in its natural state and intervening to set its value -- is the foundation of modern causal inference.

In formal notation, P(Outcome | CPAP = yes) is the probability of a good outcome among counties that happen to have CPAPs. P(Outcome | do(CPAP = yes)) is the probability of a good outcome if you give a county a CPAP. These are different quantities. They can have very different values. And confusing them is one of the most common and consequential errors in health program evaluation.

The mechanics of the do-operator are elegant. When you intervene on a variable, you perform surgery on the causal graph. You cut all the arrows pointing into the intervened variable and replace them with your decision. The variable is no longer caused by its natural parents. It is caused by you. But all the arrows pointing out of it remain. It still affects its downstream variables. It is just no longer a consequence of its upstream causes.

Consider a model where county wealth affects both CPAP availability and staff quality, and both CPAP availability and staff quality affect neonatal outcomes. When you observe that a county has a CPAP, you are also learning something about county wealth -- because wealth caused the CPAP to be there. That indirect information about wealth propagates through the model and inflates the apparent effect of the CPAP. Some of what looks like a CPAP effect is really a wealth effect.

When you perform do(CPAP = yes), you cut the arrow from wealth to CPAP. The CPAP is now set by your decision, not by county wealth. It no longer carries information about wealth. The only effect that flows from CPAP to outcomes is the genuine causal effect -- through improved respiratory care -- uncontaminated by the wealth confounder.

This is graph surgery, and it transforms a muddled observational question into a clean causal one.

Now, why does this matter for real decisions?

When a donor organization decides where to allocate funds, it needs to evaluate interventions -- things it will do. Should we buy CPAPs? Should we fund training? Should we strengthen supply chains? Each of these is an intervention, a do-operation. It changes the world.

But most of the evidence available to inform these decisions is observational. We have data on which counties have CPAPs and how they perform. We have data on which counties invested in training and what happened. We have cross-country comparisons of health spending and outcomes. All of this is observation, not intervention.

The do-operator tells you how to bridge the gap. And the bridge depends entirely on the causal model. With the right model, you can identify which confounders need to be accounted for, which variables should not be adjusted for, and whether the causal effect is even identifiable from the available data.

This brings us to influence diagrams -- an extension of causal models that explicitly incorporate decisions and objectives. An influence diagram adds decision nodes, the choices you control, and objective nodes, the things you are trying to optimize. Your decision to fund CPAPs flows through the causal model to affect your objective, say expected lives saved per dollar.

For each possible decision, you compute the expected value of your objective by propagating the intervention through the model. The decision with the highest expected value wins -- in principle. But expected value is only as good as your model.

This is where the Expected Value of Perfect Information, or EVPI, becomes useful. EVPI answers the question: how much better could you decide if you knew the true value of an uncertain variable before choosing? If EVPI for a particular variable is high, your decision is sensitive to that uncertainty, and you should invest in learning more before acting. If EVPI is low, the uncertainty does not change your decision, and you can proceed.

This is a profound idea. It tells you which uncertainties are worth resolving and which are not. In a world of limited research budgets, it is a rational guide to what to study next.

The framework of influence diagrams, do-operators, and expected value analysis transforms decision-making from guesswork into structured reasoning. It does not remove judgment. You still need to build the model, estimate the probabilities, and define your objectives. But it prevents the most common error: treating an observation as if it were an intervention.

Watching is not doing. Once you see this, every regression table looks different. Every "association" between a program and an outcome comes with a question mark. And every decision about scaling a program demands that you think about what would happen if you made it true, not just what you saw when it happened to be true.
""",
})

# ══════════════════════════════════════════════════════════════════════════════
# CHAPTER 8
# ══════════════════════════════════════════════════════════════════════════════
essays.append({
    "number": 8,
    "title": "Portfolio Thinking for Lives Saved",
    "subtitle": "Why the same logic that diversifies a stock portfolio applies to saving lives across four countries",
    "body": r"""
You have a budget. You have four countries -- Ethiopia, Rwanda, Kenya, and Tanzania. Each country has a different health system, different strengths, different gaps, and different risks. How do you allocate?

The naive approach is to rank countries by expected cost-effectiveness and pour everything into the top-ranked one. Ethiopia has the largest population. It has an established Health Extension Program. The expected cost per life saved might be the lowest. Put everything there.

This is the equivalent of putting your entire retirement savings into a single stock because it has the highest expected return. It maximizes expected performance. It also maximizes the probability of catastrophe.

Here is why.

Your estimate of cost-effectiveness for Ethiopia -- or any country -- is not a known quantity. It is a guess wrapped in uncertainty. Your model says the cost per life saved is roughly $3,000. But the confidence interval might stretch from $1,500 to $12,000. You do not know whether the government will sustain its health spending. You do not know whether the supply chain will function. You do not know whether a political crisis, a drought, or a pandemic will disrupt everything.

If you put everything into Ethiopia and the cost turns out to be $12,000 instead of $3,000, you have spent your entire budget for a quarter of the impact you expected. And you have no backup.

In 1952, an economist named Harry Markowitz won the Nobel Prize for an idea that seems obvious in retrospect: do not evaluate investments in isolation. The risk of a portfolio depends not just on the risk of each individual asset but on how those assets move relative to each other. Two risky investments that are uncorrelated can produce a portfolio that is less risky than either one alone.

This idea applies directly to health resource allocation.

Ethiopia's risks are partly driven by its unique political dynamics and its vast geography. Rwanda's risks are different -- it is a small, tightly governed country where the main uncertainty might be whether the government's current commitment to health will persist through leadership transitions. Kenya's risks come from devolution -- 47 counties making independent decisions, creating enormous variation. Tanzania's risks relate to its health information systems and supply chain capacity.

These risks are largely uncorrelated. A political crisis in Ethiopia does not cause a supply chain collapse in Tanzania. A change in Rwandan leadership does not affect Kenya's county budgets. When you spread your investment across all four, a bad outcome in one country is likely to be offset by a normal or good outcome in another.

This is diversification. It does not increase your expected return. But it dramatically reduces the chance of catastrophic underperformance.

Portfolio thinking goes further than simple diversification, though. It asks: what is the optimal allocation? Not equal shares. Not proportional to population. The optimal mix depends on each country's expected return, each country's risk, and the correlations between them.

A country with high expected return and high risk might get a moderate allocation -- you want the upside but need to hedge the downside. A country with moderate expected return and low risk might get a larger allocation than you would expect, because it provides a reliable base. A country whose risks are negatively correlated with others might get an allocation boost because it acts as insurance for the portfolio.

The tool for exploring this is Monte Carlo simulation. Instead of computing a single expected outcome for each allocation strategy, you run thousands of scenarios. In each scenario, you randomly draw values for every uncertain variable -- government co-financing rates, supply chain reliability, staff retention, community uptake -- and compute the resulting lives saved. After thousands of draws, you have a full distribution of outcomes for each strategy.

This lets you ask questions that expected-value analysis alone cannot answer. What is the probability that the portfolio saves fewer than 50,000 lives? What is the worst five-percent scenario? What allocation minimizes the chance of catastrophic failure while keeping the expected outcome high?

Consider three strategies. Strategy One: concentrate everything in Ethiopia. Strategy Two: spread equally across all four countries. Strategy Three: a tiered approach -- invest most in the two countries with the highest expected return, but maintain meaningful investments in the other two as hedges.

Monte Carlo simulation might show that Strategy One has the highest expected lives saved but also the widest spread -- including scenarios where the outcome is devastating. Strategy Two has a lower expected value but a much tighter distribution. Strategy Three hits a sweet spot: nearly as high an expected value as concentration, with nearly as narrow a distribution as equal spreading.

The right answer depends on your risk tolerance. A small foundation making a one-time bet might accept concentration risk. A large multilateral donor accountable to many stakeholders needs robustness.

There is another dimension that does not exist in financial portfolios: sequential decision-making. You do not have to allocate your entire budget at once. You can commit initial funding to all four countries, observe first-year results, and then scale up in countries that are performing well while holding steady or redirecting in countries that are not.

This is an "invest, observe, scale" approach. Early investments are partly information-gathering exercises. You are buying options, not just assets.

The value of this sequential approach depends on how informative early results are. If first-year data is highly predictive of long-term outcomes, the option value is high -- you can redirect resources early. If first-year data is noisy, the option value is low -- you cannot learn much from early signals, so you might as well commit upfront.

This is Expected Value of Information applied at the portfolio level. How much is it worth to wait and learn before going all-in?

And then there is a wrinkle that financial portfolios do not have at all: the countries respond to your investment. Your money is not being placed in a passive instrument. It is entering a system with governments that have strategies of their own. We will tackle that in the next chapter. But the existence of strategic responses makes portfolio thinking even more important. You need the diversification not just because your estimates might be wrong, but because the environment itself changes in response to your decisions.

The deeper lesson of portfolio thinking is philosophical. It says: admit that you do not know which country will perform best. Admit that your models are uncertain. And then make that uncertainty work for you by building a portfolio that is robust to being wrong about any single bet.

Diversification is not about admitting defeat. It is about acknowledging reality. And acknowledging reality is the beginning of doing something useful with it.
""",
})

# ══════════════════════════════════════════════════════════════════════════════
# CHAPTER 9
# ══════════════════════════════════════════════════════════════════════════════
essays.append({
    "number": 9,
    "title": "When the World Pushes Back",
    "subtitle": "Why governments respond strategically to donor investments -- and how to design for it",
    "body": r"""
In 2013, Kenya did something that most countries only talk about. It devolved its health system to 47 county governments. Overnight, health budgets, staffing decisions, and facility management moved from Nairobi to county capitals spread across the country. It was one of the boldest governance reforms in sub-Saharan Africa.

Here is what makes the story interesting from a strategic perspective. When international donors send money to improve maternal and newborn health in Kenya, they are no longer dealing with one government. They are dealing with 47.

Each county government has its own budget, its own priorities, and its own political calculations. And each one is watching what the others do -- and what the donors do -- before deciding how to spend its own money.

This is a game. Not the kind with winners and losers, but the kind where every player's best move depends on what the other players choose. Game theory gives us the tools to understand why rational county governments sometimes make choices that produce collectively terrible health outcomes -- and how to design funding agreements that prevent it.

The health funding landscape is a strategic interaction. Donors, national governments, and county governments each make decisions about how much to invest, where to invest, and under what conditions. Each side's optimal behavior depends on what the other sides do. And the outcome depends on all their strategies together.

Let us strip this down to the simplest possible game. A donor and a county government each choose High or Low investment in health. If both invest High, the county gets excellent health outcomes. If both invest Low, outcomes are poor. If one invests High and the other invests Low, outcomes are moderate.

What is the county government's best response?

If the donor invests High, the county is tempted to invest Low. The donor is already doing the heavy lifting. Why spend scarce county money on health when the donor will spend it for you? The county gets moderate outcomes for minimal cost -- and can redirect the savings to roads, administration, or prestige projects. Not as good as full cooperation, but a pretty good deal from the county's perspective.

If the donor invests Low, the county's best move is also Low. The donor is not contributing much, so the county would have to bear the entire cost alone for only marginally better outcomes.

In both scenarios, the county's rational move is Low. This is the free-rider problem. Each player has an incentive to let the other bear the cost. The result is underinvestment by everyone -- even when all parties would prefer mutual cooperation.

Kenya's national government recognized this trap. Its response was conditional grants.

Conditional grants are funds that flow from the national government to counties, but with strings attached. The money is released only when the county meets specific health targets -- staffing ratios, facility readiness scores, immunization coverage rates. These targets are verified through Kenya's DHIS2 health information system, which tracks real-time facility data across all 47 counties.

This is a textbook commitment device. It changes the payoff matrix so that free-riding becomes costly. A county that redirects health funds to other priorities does not just get worse health outcomes. It loses the next tranche of conditional grant funding. The rational response shifts from free-riding to cooperation.

But conditional grants only work if enforcement is credible. And here is where Kenya's story gets complicated. Monitoring 47 counties is harder than monitoring one national government. Some counties have strong DHIS2 reporting systems and transparent budgets. Others have weaker data infrastructure and less accountability. The conditional grant mechanism works well where verification is strong. It works poorly where it is not.

This creates a natural experiment. Counties with strong accountability systems -- where conditional grants are genuinely conditional -- tend to maintain health spending even when donor funding flows in. Counties with weaker systems -- where conditions are loosely enforced -- show the classic free-rider pattern. Health grants arrive, county health budgets quietly shrink, and net new investment is smaller than anyone intended.

Kenya's National Hospital Insurance Fund, NHIF, adds another layer. NHIF covers formal sector workers and provides some coverage for the poorest households. But its reach is uneven across counties. In counties with high NHIF enrollment, the local government has a built-in political constituency that demands health services. Cutting health spending means angry voters. In counties with low NHIF coverage, that political pressure is weaker -- and free-riding is easier.

The pattern is clear. Accountability changes incentives. When county governments face real consequences for underinvesting in health -- lost conditional grants, angry NHIF-enrolled voters, visible DHIS2 dashboards -- they cooperate. When those mechanisms are weak, they free-ride.

Consider the hold-up problem at county level. A donor funds training for dozens of community health volunteers in a rural county. The expectation is that the county government will absorb them -- provide stipends, supervision, and integration into the formal system.

But the county government faces a different calculation. Absorbing workers costs money from its own budget. And the county knows that if the donor cares enough about these workers, the donor will extend funding rather than watch the program collapse. The donor has already invested. Walking away means admitting failure. The sunk cost works in the county's favor.

So the county delays. It forms a committee. It cites budget constraints. It promises absorption next year. And the donor extends.

This is not unique to Kenya. It happens everywhere. But Kenya's devolved system multiplies the problem by 47. Each county is a separate game with its own players and its own equilibrium.

The solution is the same across all 47 games: design the agreement so that cooperation is each player's best response.

Milestone-based disbursement works. Structure the funding so that county co-investment is required at each stage. The donor commits funding, but releases it in tranches tied to county actions -- budget lines for health workers, verified through DHIS2, confirmed by independent review. Free-riding now means losing the next tranche.

Matching grants work even better. The donor matches every shilling the county spends on health, up to a ceiling. Now the county's incentive flips. Every shilling it spends buys two shillings of health services. Not investing means leaving matching funds on the table.

The central lesson is simple but powerful: do not rely on good intentions. Design the game so that rational self-interest leads to cooperative outcomes. This is not cynicism. Most county health directors genuinely want better outcomes for their people. But they operate within political and budgetary systems that create incentives. Those incentives sometimes point away from cooperation, even when the individuals involved would prefer to cooperate.

The most effective health investments are not the ones with the best technical design. They are the ones with the best strategic design -- the ones that anticipate how the 47 county players will respond and channel those responses toward cooperation.

A plan that ignores everyone else's strategy is not a plan. It is a wish.
""",
})

# ══════════════════════════════════════════════════════════════════════════════
# CHAPTER 10
# ══════════════════════════════════════════════════════════════════════════════
essays.append({
    "number": 10,
    "title": "Can Data Discover Causes?",
    "subtitle": "Why you need both algorithms and domain expertise -- and why neither alone is enough",
    "body": r"""
Tanzania's District Health Information System -- DHIS2 -- contains millions of records. Every health facility in the country reports data on births, deaths, complications, staffing, supplies, and referrals. The data flows from dispensaries in remote villages up through district offices to the national level. It is one of the most comprehensive health databases in sub-Saharan Africa.

Here is the fantasy: feed all of this data into an algorithm, and it will tell you what causes neonatal mortality. It will discover the causal structure of the health system -- which variables drive which outcomes, through which pathways -- without needing any human expertise at all.

This fantasy is not entirely wrong. And that is exactly what makes it dangerous.

Over the past two decades, researchers have developed algorithms that can learn causal structure from observational data. These algorithms -- the PC algorithm, GES, hill-climbing search, and others -- analyze patterns of conditional independence in data to infer which variables cause which. They are real. They work. And they are genuinely impressive.

They are also limited in ways that matter enormously for practical decision-making.

Let me explain both the promise and the boundary, because understanding what data can and cannot tell you about causation is one of the most important things you will take from this course.

Start with the promise. You have two variables, A and B, and they are correlated. There are three possible explanations: A causes B, B causes A, or some third variable C causes both. In many cases, you can figure out which explanation is correct by looking at how A and B relate to other variables in the data.

Here is the intuition. If A causes B causes C -- a chain -- then A and C are correlated, but they become independent when you hold B constant. If B causes both A and C -- a fork -- you get the same pattern: correlation that disappears when you condition on B. But if A and C both cause B -- a collider -- the pattern flips: A and C are independent, but they become correlated when you condition on B.

This asymmetry between colliders and non-colliders is the key. By systematically testing conditional independence relationships across all groups of variables, an algorithm can identify which causal structures are consistent with the data and which are not.

The PC algorithm does exactly this with Tanzania's DHIS2 data, or any data set. It starts by assuming every variable might be connected to every other, then removes connections that correspond to conditional independences. It then orients edges based on collider detection. The result is a partially oriented causal graph -- a picture of the causal structure that is consistent with the observed data.

Partially oriented is the crucial phrase. And here is where the limitation starts.

Consider three variables -- A, B, and C -- connected in a chain. Whether A causes B causes C, or C causes B causes A, or B causes both A and C as a fork, the conditional independence patterns are identical. These three structures are what we call Markov equivalent. They encode the same statistical relationships and therefore cannot be distinguished by any algorithm using observational data alone.

This is not a limitation of current technology. It is a mathematical theorem. No amount of data and no future algorithm can tell these structures apart from observational data alone. The information simply is not there.

What the data can identify are colliders -- structures where two causes converge on a common effect. These have a unique statistical signature. But for chains and forks, the direction of causation remains ambiguous.

This means that running a causal discovery algorithm on Tanzania's DHIS2 data gives you a partial answer. It tells you which variables are causally related. It orients some of the arrows -- the ones involved in collider structures. But for many arrows, it says: these variables are connected, but we cannot tell which direction from data alone.

This is where domain expertise is not a luxury. It is essential.

A health systems expert looking at an unoriented edge between "staff training" and "care quality" immediately knows the direction. Training causes quality, not the other way around. This knowledge comes from understanding how the world works -- temporal order, mechanisms, institutional logic -- and no amount of data can substitute for it.

The most powerful approach combines both: let the algorithm discover the skeleton of the graph, then use expert knowledge to orient the ambiguous edges. Neither alone is sufficient. The algorithm without expertise produces ambiguous graphs. The expertise without the algorithm is subject to confirmation bias, incomplete mental models, and the inability to detect subtle statistical patterns across dozens of variables.

There is another limitation worth understanding: the faithfulness assumption. Causal discovery algorithms assume that every statistical independence in the data reflects a real structural separation in the causal graph. This assumption fails when two causal pathways between variables happen to exactly cancel each other out.

In most natural systems, exact cancellation is rare. But in policy environments, it is plausible. If a government adjusts its spending to exactly offset donor contributions -- perfect crowding out -- the data might show no association between donor spending and total health spending, even though a causal link clearly exists. The algorithm would conclude, incorrectly, that donor spending and total spending are unrelated.

Beyond algorithmic discovery, there are other data-driven approaches to causal inference worth mentioning. Instrumental variables analysis, for instance, exploits a source of variation that affects the treatment but is otherwise unrelated to the outcome.

Suppose you want to know whether facility-based delivery reduces neonatal mortality. You cannot simply compare outcomes for mothers who delivered at facilities versus those who did not, because the mothers who choose facilities are systematically different. But if distance to the nearest facility affects whether a mother delivers there -- and distance has no direct effect on neonatal outcomes other than through facility access -- then distance is an instrument. You can use the variation in delivery rates caused by distance to estimate the causal effect of facility delivery.

The catch is that the assumptions are untestable. You have to argue that distance does not affect outcomes through any other channel. What about its effect on nutrition, on access to other services, on emergency transport times? If any of these pathways exist, the instrument is invalid.

This brings us full circle. Whether you use algorithmic causal discovery, instrumental variables, or any other technique, the quality of your causal conclusions depends on the quality of your causal assumptions. Data can refine, test, and sometimes overturn those assumptions. But it cannot generate them from nothing.

The right workflow -- the one this course has been building toward -- looks like this. First, build a qualitative causal model from domain expertise and interviews. Second, quantify it with probabilities. Third, use it for decisions -- intervention design, resource allocation, situational analysis. Fourth, use data to test and refine the model. Fifth, update and iterate.

This workflow treats expert knowledge and data as complements, not substitutes. The expert brings structure. The data brings calibration. The expert proposes. The data refines.

The scientists who built causal inference -- Judea Pearl, Peter Spirtes, Jamie Robins, and others -- all arrived at the same conclusion: causation is not in the data. It is in the model. Data can inform, test, and update the model. But the model must come first.

This sounds like a limitation. It is. But it is also a liberation. It means that causal analysis is not a passive exercise in data processing. It is an active exercise in thinking -- in drawing, questioning, debating, and refining your understanding of the system you are trying to change.

The data does not speak for itself. It never did. But with the right model, it can speak volumes.
""",
})


# ── Document Generation ───────────────────────────────────────────────────────

def build_document(essays, output_path):
    doc = Document()

    # -- Page setup --
    for section in doc.sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(2.54)
        section.right_margin = Cm(2.54)

    # -- Default font --
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(11)
    font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    pf = style.paragraph_format
    pf.space_after = Pt(6)
    pf.line_spacing = 1.15

    # -- Heading styles --
    for level in range(1, 4):
        h = doc.styles[f"Heading {level}"]
        h.font.name = "Calibri Light"
        h.font.color.rgb = RGBColor(0x00, 0x5C, 0xB9)  # BKA primary blue
        if level == 1:
            h.font.size = Pt(26)
            h.paragraph_format.space_before = Pt(0)
            h.paragraph_format.space_after = Pt(4)
        elif level == 2:
            h.font.size = Pt(16)
            h.font.italic = True
            h.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
            h.paragraph_format.space_before = Pt(0)
            h.paragraph_format.space_after = Pt(18)
        elif level == 3:
            h.font.size = Pt(14)
            h.paragraph_format.space_before = Pt(12)
            h.paragraph_format.space_after = Pt(6)

    # ── Title Page ─────────────────────────────────────────────────────────
    for _ in range(6):
        doc.add_paragraph()

    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_para.add_run("Inference and Intervention")
    run.font.name = "Calibri Light"
    run.font.size = Pt(36)
    run.font.color.rgb = RGBColor(0x00, 0x5C, 0xB9)
    run.bold = True

    subtitle_para = doc.add_paragraph()
    subtitle_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle_para.add_run("Pre-Class Essays")
    run.font.name = "Calibri Light"
    run.font.size = Pt(22)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    doc.add_paragraph()

    author_para = doc.add_paragraph()
    author_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = author_para.add_run("BK Advisors")
    run.font.name = "Calibri"
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x00, 0x5C, 0xB9)

    doc.add_page_break()

    # ── Table of Contents (manual) ─────────────────────────────────────────
    toc_heading = doc.add_heading("Contents", level=1)

    for essay in essays:
        toc_para = doc.add_paragraph()
        toc_para.paragraph_format.space_after = Pt(4)
        if essay["number"] == 0:
            toc_label = f"Introduction:  {essay['title']}"
        else:
            toc_label = f"Chapter {essay['number']}:  {essay['title']}"
        run = toc_para.add_run(toc_label)
        run.font.name = "Calibri"
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
        # Subtitle line
        sub_para = doc.add_paragraph()
        sub_para.paragraph_format.space_after = Pt(10)
        sub_para.paragraph_format.left_indent = Cm(1.27)
        run = sub_para.add_run(essay["subtitle"])
        run.font.name = "Calibri"
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
        run.italic = True

    doc.add_page_break()

    # ── Essays ─────────────────────────────────────────────────────────────
    for i, essay in enumerate(essays):
        # Chapter heading
        if essay["number"] == 0:
            doc.add_heading("Introduction", level=3)
        else:
            doc.add_heading(f"Chapter {essay['number']}", level=3)
        doc.add_heading(essay["title"], level=1)
        doc.add_heading(essay["subtitle"], level=2)

        # Body text — split into paragraphs
        body = essay["body"].strip()
        paragraphs = [p.strip() for p in body.split("\n\n") if p.strip()]

        for para_text in paragraphs:
            # Handle line continuations within paragraphs
            clean_text = " ".join(para_text.split("\n"))
            p = doc.add_paragraph(clean_text)
            p.paragraph_format.first_line_indent = Cm(0)
            p.paragraph_format.space_after = Pt(8)

        # Page break after each essay except the last
        if i < len(essays) - 1:
            doc.add_page_break()

    # ── Save ───────────────────────────────────────────────────────────────
    doc.save(output_path)
    print(f"Document saved to: {output_path}")
    print(f"Total essays: {len(essays)}")


if __name__ == "__main__":
    output = os.path.join(os.path.dirname(__file__), "pre-class-essays.docx")
    build_document(essays, output)
